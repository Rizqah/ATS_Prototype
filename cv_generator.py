from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.pdfgen import canvas as pdf_canvas
import io
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ats_engine import get_embedding

# ======================================================
# PAGE-NUMBERED CANVAS
# ======================================================
class NumberedCanvas(pdf_canvas.Canvas):
    """Canvas that draws page numbers at the bottom centre of each page."""

    def __init__(self, *args, **kwargs):
        pdf_canvas.Canvas.__init__(self, *args, **kwargs)
        self._page_num = 0

    def showPage(self):
        self._page_num += 1
        self.saveState()
        self.setFont("Helvetica", 9)
        self.drawCentredString(4.25 * inch, 0.45 * inch, f"Page {self._page_num}")
        self.restoreState()
        pdf_canvas.Canvas.showPage(self)


# ======================================================
# CV GENERATION ENGINE
# ======================================================

def generate_ats_optimized_cv(
    profile: dict,
    work_experiences: List[dict],
    achievements_by_experience: Dict[str, List[dict]],
    skills: List[dict],
    job_description: str = None
) -> tuple[bytes, float]:
    """
    Generate an ATS-optimized CV as PDF bytes.
    
    Args:
        profile: User profile data
        work_experiences: List of work experience objects
        achievements_by_experience: Dict mapping experience_id to achievements
        skills: List of skill objects
        job_description: Optional JD to tailor CV
        
    Returns:
        Tuple of (pdf_bytes, match_score)
    """
    
    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        pdf_buffer,
        pagesize=letter,
        rightMargin=0.5*inch,
        leftMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    # Define ATS-friendly styles: justified, 1.5 line spacing
    styles = getSampleStyleSheet()
    # 1.5 line spacing: leading = fontSize * 1.5 (in points)
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.black,
        spaceAfter=6,
        fontName='Helvetica-Bold',
        alignment=TA_JUSTIFY,
        leading=14 * 1.5,
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=colors.black,
        spaceAfter=4,
        spaceBefore=8,
        fontName='Helvetica-Bold',
        borderPadding=2,
        alignment=TA_JUSTIFY,
        leading=11 * 1.5,
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubHeading',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.black,
        spaceAfter=2,
        fontName='Helvetica-Bold',
        alignment=TA_JUSTIFY,
        leading=10 * 1.5,
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.black,
        spaceAfter=3,
        leading=9 * 1.5,  # 1.5 line spacing (13.5 pt)
        alignment=TA_JUSTIFY,
    )
    
    # Numbered list items: each starts as a new paragraph (space before)
    numbered_item_style = ParagraphStyle(
        'NumberedItem',
        parent=normal_style,
        spaceBefore=8,
        spaceAfter=2,
    )
    
    # Build document content (ATS-friendly: no tables, paragraphs only)
    content = []
    
    # ===== HEADER (Contact Info) – paragraphs only for ATS =====
    content.append(Paragraph(f"<b>{profile.get('full_name', 'Your Name')}</b>", title_style))
    contact_info = []
    if profile.get('email'):
        contact_info.append(profile['email'])
    if profile.get('phone'):
        contact_info.append(profile['phone'])
    if profile.get('location'):
        contact_info.append(profile['location'])
    if contact_info:
        content.append(Paragraph(" | ".join(contact_info), normal_style))
    content.append(Spacer(1, 0.15*inch))
    
    # ===== PROFESSIONAL SUMMARY =====
    if profile.get('professional_summary'):
        content.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        content.append(Paragraph(profile['professional_summary'], normal_style))
        content.append(Spacer(1, 0.1*inch))
    
    # ===== EXPERIENCE (numbered 1., 2., 3.) =====
    if work_experiences:
        content.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
        
        for exp_num, exp in enumerate(work_experiences, 1):
            # Numbered work experience: 1. Position | Company
            position_text = f"{exp_num}. <b>{exp.get('position', '')}</b> | {exp.get('company', '')}"
            content.append(Paragraph(position_text, subheading_style))
            
            # Dates
            start_date = exp.get('start_date', '')
            end_date = exp.get('end_date', '')
            if start_date:
                date_range = start_date
                if end_date:
                    date_range += f" – {end_date}"
                elif exp.get('current_job'):
                    date_range += " – Present"
                content.append(Paragraph(date_range, normal_style))
            
            # Description
            if exp.get('description'):
                content.append(Paragraph(exp['description'], normal_style))
            
            # Achievements (numbered list – each item a new paragraph)
            exp_id = exp.get('id')
            if exp_id and exp_id in achievements_by_experience:
                for idx, achievement in enumerate(achievements_by_experience[exp_id], 1):
                    text = f"{idx}. {achievement.get('achievement', '')}"
                    if achievement.get('metric'):
                        text += f" ({achievement['metric']})"
                    content.append(Paragraph(text, numbered_item_style))
            
            content.append(Spacer(1, 0.08*inch))
    
    # ===== SKILLS =====
    if skills:
        content.append(Paragraph("SKILLS", heading_style))
        
        # Group by proficiency
        skill_text = []
        for skill in skills:
            skill_text.append(skill.get('skill_name', ''))
        
        if skill_text:
            content.append(Paragraph(", ".join(skill_text), normal_style))
        
        content.append(Spacer(1, 0.1*inch))
    
    # Build PDF with numbered pages
    doc.build(content, canvasmaker=NumberedCanvas)
    pdf_bytes = pdf_buffer.getvalue()
    
    # Calculate match score if JD provided
    match_score = 0.0
    if job_description:
        # Build CV text for matching
        cv_content = f"{profile.get('full_name', '')} "
        cv_content += profile.get('professional_summary', '') + " "
        for exp in work_experiences:
            cv_content += f"{exp.get('position', '')} {exp.get('company', '')} "
            cv_content += exp.get('description', '') + " "
            exp_id = exp.get('id')
            if exp_id and exp_id in achievements_by_experience:
                for ach in achievements_by_experience[exp_id]:
                    cv_content += ach.get('achievement', '') + " "
        cv_content += " ".join([s.get('skill_name', '') for s in skills])
        
        # Get match score
        from sklearn.metrics.pairwise import cosine_similarity

        cv_vector = get_embedding(cv_content)
        jd_vector = get_embedding(job_description)
        
        if cv_vector and jd_vector:
            match_score = float(cosine_similarity([jd_vector], [cv_vector])[0][0])
    
    return pdf_bytes, match_score

def format_cv_for_display(
    profile: dict,
    work_experiences: List[dict],
    achievements_by_experience: Dict[str, List[dict]],
    skills: List[dict]
) -> str:
    """
    Format CV as plain text for display/storage.
    """
    text = []
    
    # Header
    text.append(profile.get('full_name', 'Your Name'))
    contact = []
    if profile.get('email'):
        contact.append(profile['email'])
    if profile.get('phone'):
        contact.append(profile['phone'])
    if profile.get('location'):
        contact.append(profile['location'])
    if contact:
        text.append(" | ".join(contact))
    
    # Summary
    if profile.get('professional_summary'):
        text.append("\nPROFESSIONAL SUMMARY")
        text.append(profile['professional_summary'])
    
    # Experience
    if work_experiences:
        text.append("\nPROFESSIONAL EXPERIENCE")
        for exp in work_experiences:
            text.append(f"\n{exp.get('position')} | {exp.get('company')}")
            start = exp.get('start_date', '')
            end = exp.get('end_date', '')
            if start:
                date_range = start
                if end:
                    date_range += f" – {end}"
                elif exp.get('current_job'):
                    date_range += " – Present"
                text.append(date_range)
            
            if exp.get('description'):
                text.append(exp['description'])
            
            exp_id = exp.get('id')
            if exp_id and exp_id in achievements_by_experience:
                for ach in achievements_by_experience[exp_id]:
                    ach_text = f"• {ach.get('achievement', '')}"
                    if ach.get('metric'):
                        ach_text += f" ({ach['metric']})"
                    text.append(ach_text)
    
    # Skills
    if skills:
        text.append("\nSKILLS")
        skill_names = [s.get('skill_name', '') for s in skills]
        text.append(", ".join(skill_names))
    
    return "\n".join(text)


# ======================================================
# DOCX CV GENERATION (Microsoft Word – professional, editable)
# ======================================================
def generate_cv_docx(
    profile: dict,
    work_experiences: List[dict],
    achievements_by_experience: Dict[str, List[dict]],
    skills: List[dict],
) -> bytes:
    """
    Generate CV as DOCX (Microsoft Word) for easy editing.
    Professional styling, numbered sections, justified body text.
    Returns:
        DOCX file as bytes.
    """
    doc = Document()
    doc.sections[0].top_margin = Inches(0.5)
    doc.sections[0].bottom_margin = Inches(0.5)
    doc.sections[0].left_margin = Inches(0.5)
    doc.sections[0].right_margin = Inches(0.5)

    # ---- Name ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(profile.get("full_name", "Your Name"))
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    # ---- Contact ----
    contact = []
    if profile.get("email"):
        contact.append(profile["email"])
    if profile.get("phone"):
        contact.append(profile["phone"])
    if profile.get("location"):
        contact.append(profile["location"])
    if contact:
        p = doc.add_paragraph(" | ".join(contact))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"
    doc.add_paragraph()  # spacing

    # ---- 1. PROFESSIONAL SUMMARY ----
    if profile.get("professional_summary"):
        p = doc.add_paragraph()
        p.add_run("1. PROFESSIONAL SUMMARY").bold = True
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        p2 = doc.add_paragraph(profile["professional_summary"])
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.line_spacing = 1.5
        for run in p2.runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        doc.add_paragraph()

    # ---- 2. PROFESSIONAL EXPERIENCE (numbered roles) ----
    if work_experiences:
        p = doc.add_paragraph()
        p.add_run("2. PROFESSIONAL EXPERIENCE").bold = True
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

        for exp_num, exp in enumerate(work_experiences, 1):
            # Role: 1. Position | Company
            role_p = doc.add_paragraph()
            role_p.add_run(f"{exp_num}. ").bold = True
            role_p.add_run(f"{exp.get('position', '')} | {exp.get('company', '')}").bold = True
            role_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            role_p.paragraph_format.space_before = Pt(6)
            role_p.paragraph_format.space_after = Pt(2)
            for run in role_p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"

            # Dates
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            if start:
                date_str = start
                if end:
                    date_str += f" – {end}"
                elif exp.get("current_job"):
                    date_str += " – Present"
                date_p = doc.add_paragraph(date_str)
                date_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                date_p.paragraph_format.space_after = Pt(2)
                for run in date_p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

            # Description
            if exp.get("description"):
                desc_p = doc.add_paragraph(exp["description"])
                desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                desc_p.paragraph_format.line_spacing = 1.5
                desc_p.paragraph_format.space_after = Pt(4)
                for run in desc_p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"

            # Achievements (numbered)
            exp_id = exp.get("id")
            if exp_id and exp_id in achievements_by_experience:
                for ach_num, ach in enumerate(achievements_by_experience[exp_id], 1):
                    ach_text = f"{ach_num}. {ach.get('achievement', '')}"
                    if ach.get("metric"):
                        ach_text += f" ({ach['metric']})"
                    ach_p = doc.add_paragraph(ach_text)
                    ach_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    ach_p.paragraph_format.space_before = Pt(4)
                    ach_p.paragraph_format.space_after = Pt(2)
                    ach_p.paragraph_format.line_spacing = 1.5
                    for run in ach_p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
        doc.add_paragraph()

    # ---- 3. SKILLS ----
    if skills:
        p = doc.add_paragraph()
        p.add_run("3. SKILLS").bold = True
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        skill_names = [s.get("skill_name", "") for s in skills]
        skills_p = doc.add_paragraph(", ".join(skill_names))
        skills_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        skills_p.paragraph_format.line_spacing = 1.5
        for run in skills_p.runs:
            run.font.size = Pt(9)
            run.font.name = "Calibri"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ======================================================
# OPTIMIZED CV CONTENT (AI-REWRITTEN FOR JD)
# ======================================================
def generate_optimized_cv_content(
    profile: dict,
    work_experiences: List[dict],
    achievements_by_experience: Dict[str, List[dict]],
    skills: List[dict],
    job_description: str
) -> dict:
    """
    Generate optimized CV content using AI.
    Expands on relevant skills, reorders achievements, rewrites professional summary.

    Returns:
        dict with profile, work_experiences, achievements_by_exp, skills, cv_text
    """
    from ats_engine import optimize_cv_for_jd

    profile_text = f"""
    Current Profile:
    {profile.get('full_name', '')}

    Summary: {profile.get('professional_summary', '')}

    Skills: {', '.join([s.get('skill_name', '') for s in skills])}
    """

    work_exp_text = ""
    for exp in work_experiences:
        work_exp_text += f"""
        Position: {exp.get('position', '')}
        Company: {exp.get('company', '')}
        Period: {exp.get('start_date', '')} to {exp.get('end_date', 'Present' if exp.get('current_job') else '')}
        Description: {exp.get('description', '')}

        Achievements:
        """
        exp_id = exp.get('id')
        if exp_id and exp_id in achievements_by_experience:
            for ach in achievements_by_experience[exp_id]:
                work_exp_text += f"- {ach.get('achievement', '')} ({ach.get('metric', '')})\n"

    optimized_content = optimize_cv_for_jd(profile_text, job_description, work_exp_text)

    optimized_profile = profile.copy()
    optimized_experiences = work_experiences.copy()
    optimized_skills = skills.copy()
    optimized_achievements = achievements_by_experience.copy()

    if "OPTIMIZED_SUMMARY:" in optimized_content:
        summary_start = optimized_content.find("OPTIMIZED_SUMMARY:") + len("OPTIMIZED_SUMMARY:")
        summary_end = optimized_content.find("\n\nOPTIMIZED_EXPERIENCE:")
        if summary_end == -1:
            summary_end = optimized_content.find("\nOPTIMIZED_SKILLS:")
        optimized_profile["professional_summary"] = optimized_content[summary_start:summary_end].strip()

    if "OPTIMIZED_SKILLS:" in optimized_content:
        skills_start = optimized_content.find("OPTIMIZED_SKILLS:") + len("OPTIMIZED_SKILLS:")
        skills_text = optimized_content[skills_start:].strip().split('\n')[0]
        optimized_skills = [{"skill_name": s.strip(), "proficiency": "Expert"} for s in skills_text.split(',')]

    cv_text = f"""
    {optimized_profile.get('full_name', '')}
    {optimized_profile.get('email', '')} | {optimized_profile.get('phone', '')} | {optimized_profile.get('location', '')}

    PROFESSIONAL SUMMARY
    {optimized_profile.get('professional_summary', '')}

    PROFESSIONAL EXPERIENCE
    """
    for idx, exp in enumerate(optimized_experiences, 1):
        cv_text += f"""
    {idx}. {exp.get('position', '')} | {exp.get('company', '')}
    {exp.get('start_date', '')} – {exp.get('end_date', 'Present' if exp.get('current_job') else '')}
    {exp.get('description', '')}
    """
        exp_id = exp.get('id')
        if exp_id and exp_id in optimized_achievements:
            for ach in optimized_achievements[exp_id]:
                cv_text += f"• {ach.get('achievement', '')} ({ach.get('metric', '')})\n"
    cv_text += "\nSKILLS\n"
    cv_text += ", ".join([s.get('skill_name', '') for s in optimized_skills])

    return {
        "profile": optimized_profile,
        "work_experiences": optimized_experiences,
        "achievements_by_exp": optimized_achievements,
        "skills": optimized_skills,
        "cv_text": cv_text
    }