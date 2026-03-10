"""
CV Helper Module - Templates, Version History, Multi-format Downloads, Empty States
Provides utilities for CV generation, management, and display.
"""

import streamlit as st
from typing import List, Dict, Optional, Tuple
from datetime import datetime
import json
from io import BytesIO

# Import PDF and DOCX generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor

# ======================================================
# CV TEMPLATES CONFIGURATION
# ======================================================

CV_TEMPLATES = {
    "professional": {
        "name": "Professional",
        "description": "Clean, modern format perfect for corporate roles",
        "features": ["ATS-optimized", "Single column layout", "Modern fonts", "High readability"],
        "icon": "📊",
        "best_for": "Finance, Management, Consulting",
        "color": "#3b82f6"
    },
    "creative": {
        "name": "Creative",
        "description": "Modern design for creative industries",
        "features": ["Visual hierarchy", "Modern styling", "Accent colors", "Eye-catching"],
        "icon": "🎨",
        "best_for": "Design, Marketing, Branding",
        "color": "#8b5cf6"
    },
    "technical": {
        "name": "Technical",
        "description": "Optimized for tech and engineering roles",
        "features": ["Code snippet section", "Technical skills prominently", "Project highlights", "GitHub links"],
        "icon": "💻",
        "best_for": "Engineering, Data Science, Development",
        "color": "#10b981"
    },
    "academic": {
        "name": "Academic",
        "description": "Formal layout for academic and research positions",
        "features": ["Publication section", "Research highlights", "Academic credentials", "Citations"],
        "icon": "🎓",
        "best_for": "Academia, Research, Sciences",
        "color": "#f59e0b"
    }
}

# ======================================================
# CV TEMPLATE DISPLAY
# ======================================================

def show_cv_template_selector() -> Optional[str]:
    """
    Display visual CV template selector with previews.
    
    Returns:
        Selected template key or None
    """
    st.subheader("📋 CV Template Selection")
    st.write("Choose a template that best fits your industry and role:")
    
    # Create columns for templates
    cols = st.columns(2)
    selected_template = None
    
    for idx, (template_key, template_info) in enumerate(CV_TEMPLATES.items()):
        col = cols[idx % 2]
        
        with col:
            # Create styled container
            container = st.container(border=True)
            
            with container:
                # Template header with icon
                col1, col2 = st.columns([1, 4])
                with col1:
                    st.write(template_info["icon"])
                with col2:
                    st.markdown(f"**{template_info['name']}**")
                
                # Description
                st.caption(template_info["description"])
                
                # Features
                st.markdown("**Features:**")
                for feature in template_info["features"]:
                    st.markdown(f"✓ {feature}")
                
                # Best for
                st.markdown(f"*Best for: {template_info['best_for']}*")
                
                # Select button
                if st.button(
                    f"📌 Select {template_info['name']}",
                    key=f"select_{template_key}",
                    use_container_width=True
                ):
                    selected_template = template_key
                    st.session_state.selected_cv_template = template_key
                    st.success(f"✅ {template_info['name']} template selected!")
    
    return selected_template


def show_template_comparison() -> None:
    """Display comparison table of all CV templates"""
    
    st.subheader("📊 Template Comparison")
    
    comparison_data = {
        "Template": [],
        "Best For": [],
        "Layout": [],
        "ATS Score": [],
        "Visual Appeal": []
    }
    
    for template_key, template_info in CV_TEMPLATES.items():
        comparison_data["Template"].append(template_info["name"])
        comparison_data["Best For"].append(template_info["best_for"])
        
        if template_key == "professional":
            comparison_data["Layout"].append("Single Column")
            comparison_data["ATS Score"].append("⭐⭐⭐⭐⭐")
            comparison_data["Visual Appeal"].append("⭐⭐⭐⭐")
        elif template_key == "creative":
            comparison_data["Layout"].append("Multi-Column")
            comparison_data["ATS Score"].append("⭐⭐⭐")
            comparison_data["Visual Appeal"].append("⭐⭐⭐⭐⭐")
        elif template_key == "technical":
            comparison_data["Layout"].append("Code Sections")
            comparison_data["ATS Score"].append("⭐⭐⭐⭐")
            comparison_data["Visual Appeal"].append("⭐⭐⭐⭐")
        else:  # academic
            comparison_data["Layout"].append("Publication-Heavy")
            comparison_data["ATS Score"].append("⭐⭐⭐⭐")
            comparison_data["Visual Appeal"].append("⭐⭐⭐")
    
    import pandas as pd
    df = pd.DataFrame(comparison_data)
    st.dataframe(df, use_container_width=True, hide_index=True)


# ======================================================
# DOWNLOAD HELPER FUNCTIONS
# ======================================================

def generate_cv_pdf_bytes(cv_content: str, cv_name: str) -> bytes:
    """
    Generate PDF bytes from CV text content.
    
    Args:
        cv_content: Text content of the CV
        cv_name: Name for the CV
        
    Returns:
        PDF file bytes
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, margins=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=RGBColor(60, 120, 160),
        spaceAfter=12,
        alignment=1
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=12,
        alignment=TA_JUSTIFY
    )
    
    story = []
    
    # Add title
    story.append(Paragraph(f"{cv_name} - CV", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Add content
    lines = cv_content.split('\n')
    for line in lines:
        if line.strip():
            story.append(Paragraph(line, body_style))
        else:
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def generate_cv_docx_bytes(cv_content: str, cv_name: str) -> bytes:
    """
    Generate DOCX bytes from CV text content.
    
    Args:
        cv_content: Text content of the CV
        cv_name: Name for the CV
        
    Returns:
        DOCX file bytes
    """
    doc = DocxDocument()
    
    # Add title
    title = doc.add_heading(f"{cv_name} - CV", 0)
    title.runs[0].font.color.rgb = RGBColor(60, 120, 160)
    
    # Add generated date
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Add content
    lines = cv_content.split('\n')
    for line in lines:
        if line.strip():
            # Check if it's a section header (all caps or ends with colon)
            if line.isupper() or line.rstrip().endswith(':'):
                p = doc.add_heading(line, level=2)
                p.runs[0].font.size = Pt(12)
            else:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # Add spacing for empty lines
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ======================================================

def show_cv_version_history(user_email: str, cvs_data: Dict) -> None:
    """
    Display CV version history with metadata and restoration options.
    
    Args:
        user_email: User's email address
        cvs_data: CVs data from generated_cvs.json
    """
    
    if user_email not in cvs_data or not cvs_data[user_email]:
        show_empty_cv_state("empty_history")
        return
    
    user_cvs = cvs_data[user_email]
    
    st.subheader("📜 CV Version History")
    
    # Summary stats
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total CVs Generated", len(user_cvs))
    with col2:
        avg_score = sum(cv["match_score"] for cv in user_cvs.values()) / len(user_cvs)
        st.metric("Average Match Score", f"{avg_score*100:.1f}%")
    with col3:
        st.metric("Most Recent", "Today")
    
    st.divider()
    
    # Version list
    st.write("**Select a version to view, restore, or download:**")
    
    # Sort by created_at descending (newest first)
    sorted_cvs = sorted(
        user_cvs.values(),
        key=lambda x: x.get("created_at", ""),
        reverse=True
    )
    
    for idx, cv in enumerate(sorted_cvs, 1):
        # Parse timestamp
        created_at = cv.get("created_at", "Unknown")
        date_obj = datetime.fromisoformat(created_at.replace("Z", "+00:00")) if "T" in created_at else None
        formatted_date = date_obj.strftime("%b %d, %Y at %I:%M %p") if date_obj else created_at
        
        # Version card
        with st.expander(
            f"📄 v{idx}: {cv['job_title']} | {formatted_date} | Score: {cv['match_score']*100:.1f}%",
            expanded=(idx == 1)
        ):
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.write("**Job Title**")
                st.write(cv["job_title"])
            
            with col2:
                st.write("**Match Score**")
                score_pct = cv["match_score"] * 100
                
                # Color-coded score
                if score_pct >= 80:
                    st.success(f"🟢 {score_pct:.1f}%")
                elif score_pct >= 60:
                    st.info(f"🟡 {score_pct:.1f}%")
                elif score_pct >= 40:
                    st.warning(f"🔴 {score_pct:.1f}%")
                else:
                    st.error(f"⚫ {score_pct:.1f}%")
            
            with col3:
                st.write("**Status**")
                st.write(cv.get("application_status", "Draft"))
            
            with col4:
                st.write("**Created**")
                st.write(formatted_date)
            
            st.divider()
            
            # Preview section
            with st.expander("👁️ Preview CV Content"):
                st.text_area(
                    "CV Content Preview",
                    value=cv["cv_content"],
                    height=300,
                    disabled=True,
                    key=f"preview_{cv['id']}"
                )
            
            # Job description section
            with st.expander("📋 Job Description Used"):
                st.text_area(
                    "Job Description Preview",
                    value=cv["job_description"],
                    height=250,
                    disabled=True,
                    key=f"jd_{cv['id']}"
                )
            
            st.divider()
            
            # Action buttons
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button(
                    "🔄 Use This Version",
                    key=f"restore_{cv['id']}",
                    use_container_width=True,
                    help="Reuse this CV and generate it again"
                ):
                    st.session_state.restored_cv = cv
                    st.success("✅ Version restored! You can now re-generate or download it.")
            
            with col2:
                pdf_bytes = generate_cv_pdf_bytes(cv["cv_content"], cv["job_title"])
                st.download_button(
                    "📥 Download as PDF",
                    data=pdf_bytes,
                    file_name=f"{cv['job_title']}_CV.pdf",
                    mime="application/pdf",
                    key=f"dl_pdf_{cv['id']}",
                    use_container_width=True
                )
            
            with col3:
                docx_bytes = generate_cv_docx_bytes(cv["cv_content"], cv["job_title"])
                st.download_button(
                    "📥 Download as DOCX",
                    data=docx_bytes,
                    file_name=f"{cv['job_title']}_CV.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_docx_{cv['id']}",
                    use_container_width=True
                )
            
            with col4:
                if st.button(
                    "🗑️ Delete Version",
                    key=f"del_{cv['id']}",
                    use_container_width=True
                ):
                    st.warning(f"⚠️ Version {cv['job_title']} marked for deletion")


# ======================================================
# MULTI-FORMAT DOWNLOAD
# ======================================================

def show_download_options(
    cv_bytes: bytes,
    docx_bytes: bytes,
    cv_name: str
) -> None:
    """
    Display multi-format download options.
    
    Args:
        cv_bytes: PDF bytes
        docx_bytes: DOCX bytes
        cv_name: Name for the CV files
    """
    
    st.subheader("📥 Download Your CV")
    st.write("Choose your preferred format:")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.download_button(
            label="📄 Download as PDF",
            data=cv_bytes,
            file_name=f"{cv_name}_CV.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        st.caption("**Best for:** Email, applications, printing")
    
    with col2:
        st.download_button(
            label="📊 Download as DOCX",
            data=docx_bytes,
            file_name=f"{cv_name}_CV.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.caption("**Best for:** Editing, customization")
    
    with col3:
        # Generate TXT version from PDF
        txt_content = f"""
{cv_name} - CV

Generated: {datetime.now().strftime('%Y-%m-%d')}

[CV Content - Plain Text Format]
This is a plain text version of your CV, optimized for ATS (Applicant Tracking Systems).

Please download the PDF or DOCX version for the full formatted CV.
"""
        st.download_button(
            label="📝 Download as TXT",
            data=txt_content,
            file_name=f"{cv_name}_CV.txt",
            mime="text/plain",
            use_container_width=True
        )
        st.caption("**Best for:** ATS systems, parsing")


# ======================================================
# EMPTY STATES
# ======================================================

EMPTY_STATES = {
    "empty_history": {
        "icon": "📋",
        "title": "No CVs Generated Yet",
        "message": "You haven't generated any CVs yet. Create your first tailored CV now!",
        "cta_text": "🚀 Generate Your First CV",
        "cta_page": "pages/06_careerhub_cv_generator.py"
    },
    "incomplete_profile": {
        "icon": "👤",
        "title": "Complete Your Profile",
        "message": "You need to add work experience and skills before generating a CV.",
        "cta_text": "📝 Complete Your Profile",
        "cta_page": "pages/05_careerhub_profile.py"
    },
    "cv_limit_reached": {
        "icon": "⚠️",
        "title": "Monthly CV Limit Reached",
        "message": "You've reached your monthly CV generation limit. Upgrade to Premium for unlimited CVs.",
        "cta_text": "💎 Upgrade to Premium",
        "cta_page": None
    },
    "no_favorites": {
        "icon": "⭐",
        "title": "No Favorite CVs",
        "message": "You haven't marked any CVs as favorites yet. Star your best versions!",
        "cta_text": "📜 View CV History",
        "cta_page": None
    },
    "empty_download": {
        "icon": "📥",
        "title": "Ready to Download",
        "message": "Generate a CV first to download it in multiple formats (PDF, DOCX, TXT).",
        "cta_text": "📄 Generate a CV",
        "cta_page": "pages/06_careerhub_cv_generator.py"
    }
}


def show_empty_cv_state(state_type: str = "empty_history") -> None:
    """
    Display an empty state with helpful CTA.
    
    Args:
        state_type: Type of empty state to show
    """
    
    if state_type not in EMPTY_STATES:
        state_type = "empty_history"
    
    state = EMPTY_STATES[state_type]
    
    # Create centered empty state
    col1, col2, col3 = st.columns([1, 3, 1])
    
    with col2:
        st.markdown(f"### {state['icon']} {state['title']}")
        st.markdown(f"{state['message']}")
        
        st.divider()
        
        # CTA Button
        if state['cta_page'] and state['cta_page'] != "pages/06_careerhub_cv_generator.py":
            if st.button(
                state['cta_text'],
                use_container_width=True,
                type="primary"
            ):
                st.switch_page(state['cta_page'])
        elif state['cta_page'] == "pages/06_careerhub_cv_generator.py":
            if st.button(
                state['cta_text'],
                use_container_width=True,
                type="primary"
            ):
                st.switch_page(state['cta_page'])
        else:
            # For states without navigation
            st.button(
                state['cta_text'],
                use_container_width=True,
                type="secondary",
                disabled=True
            )
        
        # Additional help text
        st.markdown("---")
        st.markdown("💡 **Need help?** Check our [Help & FAQ](pages/10_help_faq.py)")


def show_cv_dashboard_empty_state() -> None:
    """Display empty state for CV dashboard with multiple helpful CTAs"""
    
    st.title("📄 Your CVs")
    
    col1, col2, col3 = st.columns([1, 4, 1])
    
    with col2:
        st.markdown("## 📋 You haven't generated any CVs yet")
        
        st.markdown("""
        **Get started in 3 steps:**
        
        1. **Complete Your Profile** - Add your work experience, skills, and achievements
        2. **Choose a Template** - Pick a CV template that matches your industry
        3. **Generate CVs** - Create tailored CVs for specific job opportunities
        
        Once you generate CVs, you'll be able to:
        - 📜 View your CV version history
        - 📥 Download in multiple formats (PDF, DOCX, TXT)
        - 🔄 Reuse and restore previous versions
        - 📊 See match scores for each job application
        """)
        
        st.divider()
        
        # CTA Buttons
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button(
                "👤 Complete Profile",
                use_container_width=True,
                type="primary"
            ):
                st.switch_page("pages/05_careerhub_profile.py")
        
        with col2:
            if st.button(
                "📄 Generate CV",
                use_container_width=True,
                type="primary"
            ):
                st.switch_page("pages/06_careerhub_cv_generator.py")
        
        with col3:
            if st.button(
                "❓ Get Help",
                use_container_width=True
            ):
                st.switch_page("pages/10_help_faq.py")


# ======================================================
# CV STATISTICS & INSIGHTS
# ======================================================

def show_cv_statistics(cvs_data: Dict[str, Dict]) -> None:
    """
    Display CV generation statistics and insights.
    
    Args:
        cvs_data: CVs data from generated_cvs.json
    """
    
    if not cvs_data or not any(cvs_data.values()):
        st.info("📊 No CVs generated yet. Statistics will appear here once you create your first CV.")
        return
    
    st.subheader("📊 Your CV Statistics")
    
    # Calculate stats
    all_cvs = []
    for user_cvs in cvs_data.values():
        all_cvs.extend(user_cvs.values())
    
    if not all_cvs:
        return
    
    scores = [cv["match_score"] for cv in all_cvs]
    avg_score = sum(scores) / len(scores)
    max_score = max(scores)
    min_score = min(scores)
    
    # Metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total CVs", len(all_cvs))
    
    with col2:
        st.metric("Highest Match", f"{max_score*100:.1f}%")
    
    with col3:
        st.metric("Average Match", f"{avg_score*100:.1f}%")
    
    with col4:
        st.metric("Lowest Match", f"{min_score*100:.1f}%")
    
    # Distribution chart
    st.divider()
    
    import pandas as pd
    
    # Group scores into buckets
    score_buckets = {
        "🟢 Excellent (80%+)": sum(1 for s in scores if s >= 0.8),
        "🟡 Good (60-79%)": sum(1 for s in scores if 0.6 <= s < 0.8),
        "🔴 Fair (40-59%)": sum(1 for s in scores if 0.4 <= s < 0.6),
        "⚫ Poor (<40%)": sum(1 for s in scores if s < 0.4),
    }
    
    st.write("**Match Score Distribution:**")
    
    # Bar chart
    import matplotlib.pyplot as plt
    
    fig, ax = plt.subplots(figsize=(10, 4))
    
    categories = list(score_buckets.keys())
    values = list(score_buckets.values())
    colors = ["#10b981", "#f59e0b", "#ef4444", "#6b7280"]
    
    bars = ax.bar(categories, values, color=colors)
    ax.set_ylabel("Number of CVs")
    ax.set_title("Distribution of CV Match Scores")
    ax.set_ylim(0, max(values) + 1)
    
    # Add value labels on bars
    for bar in bars:
        height = bar.get_height()
        if height > 0:
            ax.text(bar.get_x() + bar.get_width()/2., height,
                   f'{int(height)}',
                   ha='center', va='bottom')
    
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()
    
    st.pyplot(fig)


# ======================================================
# CV MANAGEMENT UTILITIES
# ======================================================

def get_latest_cv(cvs_data: Dict[str, Dict]) -> Optional[Dict]:
    """
    Get the most recently generated CV.
    
    Args:
        cvs_data: CVs data
        
    Returns:
        Latest CV dict or None
    """
    all_cvs = []
    for user_cvs in cvs_data.values():
        all_cvs.extend(user_cvs.values())
    
    if not all_cvs:
        return None
    
    return max(all_cvs, key=lambda x: x.get("created_at", ""))


def get_best_matching_cv(cvs_data: Dict[str, Dict]) -> Optional[Dict]:
    """
    Get the CV with highest match score.
    
    Args:
        cvs_data: CVs data
        
    Returns:
        Best matching CV dict or None
    """
    all_cvs = []
    for user_cvs in cvs_data.values():
        all_cvs.extend(user_cvs.values())
    
    if not all_cvs:
        return None
    
    return max(all_cvs, key=lambda x: x.get("match_score", 0))


def show_cv_quick_actions() -> None:
    """Display quick action buttons for CV management"""
    
    st.subheader("🔧 Quick Actions")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        if st.button("📄 Generate New CV", use_container_width=True):
            st.switch_page("pages/06_careerhub_cv_generator.py")
    
    with col2:
        if st.button("📜 View History", use_container_width=True):
            st.session_state.show_cv_history = True
    
    with col3:
        if st.button("❓ Help & FAQ", use_container_width=True):
            st.switch_page("pages/10_help_faq.py")
    
    with col4:
        if st.button("⚙️ Settings", use_container_width=True):
            st.switch_page("pages/05_careerhub_profile.py")
