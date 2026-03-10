import streamlit as st
import sys
import os

# ======================================================
# PATH SETUP
# ======================================================
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, BASE_DIR)

# ======================================================
# IMPORTS
# ======================================================
from careerhub_db import (
    save_generated_cv,
    check_cv_limit
)
from cv_generator import (
    generate_ats_optimized_cv,
    generate_optimized_cv_content,
    format_cv_for_display
)
from ats_engine import match_profile_to_jd
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from ui_components import show_top_navigation

# ======================================================
# PAGE CONFIG
# ======================================================
st.set_page_config(
    page_title="TrueFit - Optimize CV",
    page_icon="✨",
    layout="wide"
)

# ======================================================
# AUTH CHECK
# ======================================================
if "authenticated" not in st.session_state or not st.session_state.authenticated:
    st.error("❌ Please log in first")
    st.switch_page("pages/03_careerhub_auth.py")

user_email = st.session_state.user_email
show_top_navigation("CV Match")

# ======================================================
# CHECK REQUIRED DATA
# ======================================================
if "profile_data" not in st.session_state or "job_description" not in st.session_state:
    st.error("❌ No job or profile data found. Please check your CV match first.")
    st.switch_page("pages/08_careerhub_cv_matcher.py")

profile_data = st.session_state.profile_data
job_description = st.session_state.job_description
profile = profile_data["profile"]

# ======================================================
# HEADER
# ======================================================
st.title("✨ Optimize Your CV")
st.caption("Rewriting your CV to better match the job requirements")

# ======================================================
# CV LIMIT CHECK
# ======================================================
limit_data = check_cv_limit(user_email)

if limit_data.get("limit_reached"):
    st.error(
        f"❌ Monthly CV limit reached ({limit_data['limit']}).\n\n"
        "Upgrade to Premium for unlimited CVs."
    )
    st.stop()

remaining = limit_data["limit"] - limit_data["cvs_generated"]
st.info(f"📊 CVs remaining this month: {remaining}")

# ======================================================
# GENERATE OPTIMIZED CV
# ======================================================
st.header("1️⃣ Optimizing Your CV")

if st.button("🚀 Generate Optimized CV", type="primary", use_container_width=True):
    with st.spinner("🤖 Optimizing your CV based on job requirements..."):
        try:
            # Generate optimized CV content
            optimized_cv_data = generate_optimized_cv_content(
                profile=profile,
                work_experiences=profile_data["work_experiences"],
                achievements_by_experience=profile_data["achievements_by_exp"],
                skills=profile_data["skills"],
                job_description=job_description
            )

            # Generate PDF
            pdf_bytes, new_score = generate_ats_optimized_cv(
                profile=optimized_cv_data["profile"],
                work_experiences=optimized_cv_data["work_experiences"],
                achievements_by_experience=optimized_cv_data["achievements_by_exp"],
                skills=optimized_cv_data["skills"],
                job_description=job_description
            )

            # Generate DOCX
            docx_bytes = generate_docx_cv(optimized_cv_data)

            # Store in session
            st.session_state.optimized_pdf = pdf_bytes
            st.session_state.optimized_docx = docx_bytes
            st.session_state.optimized_score = new_score
            st.session_state.optimized_cv_text = optimized_cv_data["cv_text"]

            st.success("✅ CV optimized successfully!")
            st.rerun()

        except Exception as e:
            st.error(f"❌ Error optimizing CV: {str(e)}")

# ======================================================
# RESULTS & COMPARISON
# ======================================================
if "optimized_score" in st.session_state:
    st.divider()
    st.subheader("📊 Score Comparison")

    original_score = st.session_state.match_score
    optimized_score = st.session_state.optimized_score
    improvement = (optimized_score - original_score) * 100

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("Original Score", f"{original_score*100:.1f}%")

    with col2:
        if improvement >= 0:
            st.metric("Optimized Score", f"{optimized_score*100:.1f}%", f"+{improvement:.1f}%")
        else:
            st.metric("Optimized Score", f"{optimized_score*100:.1f}%", f"{improvement:.1f}%")

    with col3:
        if optimized_score >= 0.7:
            status = "Strong Fit 🟢"
        elif optimized_score >= 0.5:
            status = "Moderate Fit 🟡"
        else:
            status = "Weak Fit 🔴"
        st.metric("Status", status)

    # ===== DOWNLOAD OPTIONS =====
    st.divider()
    st.subheader("📥 Download Your CVs")

    st.write("**Original CV:**")
    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 PDF",
            data=st.session_state.generated_pdf,
            file_name=f"{profile.get('full_name','CV').replace(' ','_')}_ORIGINAL.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    with col2:
        if "generated_docx" in st.session_state:
            st.download_button(
                label="📘 DOCX",
                data=st.session_state.generated_docx,
                file_name=f"{profile.get('full_name','CV').replace(' ','_')}_ORIGINAL.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    st.write("**Optimized CV:**")
    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            label="📄 PDF",
            data=st.session_state.optimized_pdf,
            file_name=f"{profile.get('full_name','CV').replace(' ','_')}_OPTIMIZED.pdf",
            mime="application/pdf",
            use_container_width=True
        )

    with col2:
        st.download_button(
            label="📘 DOCX",
            data=st.session_state.optimized_docx,
            file_name=f"{profile.get('full_name','CV').replace(' ','_')}_OPTIMIZED.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

    # Save to database
    st.divider()
    if st.button("💾 Save CVs to My Profile", type="primary", use_container_width=True):
        try:
            save_generated_cv(
                user_email=user_email,
                job_title="Optimized CV",
                match_score=optimized_score,
                cv_content=st.session_state.optimized_cv_text,
                job_description=job_description
            )
            st.success("✅ CVs saved to your profile!")
        except Exception as e:
            st.error(f"❌ Error saving CVs: {str(e)}")

# ======================================================
# NAVIGATION
# ======================================================
st.divider()

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("⬅️ Back to Match Check", use_container_width=True):
        st.switch_page("pages/08_careerhub_cv_matcher.py")

with col2:
    if st.button("📊 View All CVs", use_container_width=True):
        st.switch_page("pages/07_careerhub_tracker.py")

with col3:
    if st.button("🎯 Check Another Job", use_container_width=True):
        for key in ["match_score", "profile_data", "job_description", "optimized_score", "optimized_pdf"]:
            if key in st.session_state:
                del st.session_state[key]
        st.switch_page("pages/08_careerhub_cv_matcher.py")

# ======================================================
# DOCX GENERATOR
# ======================================================
def generate_docx_cv(cv_data):
    """Generate CV as DOCX format."""
    doc = Document()
    
    profile = cv_data["profile"]
    work_experiences = cv_data["work_experiences"]
    achievements_by_exp = cv_data["achievements_by_exp"]
    skills = cv_data["skills"]
    
    # Header - Name
    heading = doc.add_paragraph()
    heading_run = heading.add_run(profile.get("full_name", "Your Name"))
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    
    # Contact Info
    contact = []
    if profile.get("email"):
        contact.append(profile["email"])
    if profile.get("phone"):
        contact.append(profile["phone"])
    if profile.get("location"):
        contact.append(profile["location"])
    
    if contact:
        contact_p = doc.add_paragraph(" | ".join(contact))
        contact_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Professional Summary
    if profile.get("professional_summary"):
        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(profile["professional_summary"])
    
    # Work Experience
    if work_experiences:
        doc.add_heading("Professional Experience", level=2)
        
        for idx, exp in enumerate(work_experiences, 1):
            # Position and Company
            exp_heading = doc.add_paragraph()
            exp_run = exp_heading.add_run(f"{idx}. {exp.get('position')} | {exp.get('company')}")
            exp_run.bold = True
            
            # Dates
            if exp.get("start_date"):
                date_str = exp["start_date"]
                if exp.get("end_date"):
                    date_str += f" – {exp['end_date']}"
                elif exp.get("current_job"):
                    date_str += " – Present"
                doc.add_paragraph(date_str, style="List Bullet")
            
            # Description
            if exp.get("description"):
                doc.add_paragraph(exp["description"])
            
            # Achievements
            exp_id = exp.get("id")
            if exp_id and exp_id in achievements_by_exp:
                for ach in achievements_by_exp[exp_id]:
                    ach_text = f"{ach.get('achievement')}"
                    if ach.get("metric"):
                        ach_text += f" ({ach['metric']})"
                    doc.add_paragraph(ach_text, style="List Bullet")
    
    # Skills
    if skills:
        doc.add_heading("Skills", level=2)
        skills_text = ", ".join([s.get("skill_name") for s in skills])
        doc.add_paragraph(skills_text)
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()
